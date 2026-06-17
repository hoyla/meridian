"""Tests for the markdown → docx translator (`briefing_pack.md_to_docx`).

Focused unit coverage that doesn't need a DB. Each test gives the
translator a small markdown fragment and asserts the resulting
Document has the expected structure.
"""

from __future__ import annotations

from docx import Document

from briefing_pack.md_to_docx import MarkdownToDocxTranslator


def _translate(md: str, **kwargs) -> Document:
    doc = Document()
    MarkdownToDocxTranslator(doc, **kwargs).translate(md)
    return doc


def _styles(doc: Document) -> list[tuple[str, str]]:
    return [
        (p.style.name if p.style else "?", p.text)
        for p in doc.paragraphs
    ]


def _images(doc: Document) -> int:
    return sum(
        1 for p in doc.paragraphs
        for r in p.runs
        if r.element.xpath(".//w:drawing")
    )


# ---------------------------------------------------------------------------
# Headings
# ---------------------------------------------------------------------------

class TestHeadings:
    def test_h1_through_h6(self):
        md = "# h1\n\n## h2\n\n### h3\n\n#### h4\n\n##### h5\n\n###### h6\n"
        doc = _translate(md)
        styles = [p.style.name for p in doc.paragraphs if p.text]
        # mistune attrs.level maps to docx Heading level 1-6
        assert styles == [
            "Heading 1", "Heading 2", "Heading 3",
            "Heading 4", "Heading 5", "Heading 6",
        ]

    def test_heading_with_inline_bold(self):
        doc = _translate("## Section **bold** trailing\n")
        h = next(p for p in doc.paragraphs if p.style.name == "Heading 2")
        assert h.text == "Section bold trailing"
        # Find the bold run
        bold_runs = [r for r in h.runs if r.bold]
        assert any("bold" in r.text for r in bold_runs)


# ---------------------------------------------------------------------------
# Paragraphs + inline runs
# ---------------------------------------------------------------------------

class TestInlineRuns:
    def test_paragraph_with_mixed_inline(self):
        md = "Plain **bold** and *italic* and `code` text.\n"
        doc = _translate(md)
        p = doc.paragraphs[0]
        assert p.style.name == "Normal"
        runs_by_type = {
            "bold": [r for r in p.runs if r.bold],
            "italic": [r for r in p.runs if r.italic],
        }
        assert any(r.text == "bold" for r in runs_by_type["bold"])
        assert any(r.text == "italic" for r in runs_by_type["italic"])
        # codespan run uses Courier New
        code_runs = [r for r in p.runs if r.font.name == "Courier New"]
        assert any(r.text == "code" for r in code_runs)

    def test_strong_inside_emphasis(self):
        """Bold inside italic → run with bold + italic both set."""
        md = "*outer **bold** inner* trailing\n"
        doc = _translate(md)
        p = doc.paragraphs[0]
        # Find the run that is both bold and italic
        combo = [r for r in p.runs if r.bold and r.italic]
        assert any(r.text == "bold" for r in combo)

    def test_hyperlink_creates_docx_hyperlink_element(self):
        md = "Click [here](https://example.com) please.\n"
        doc = _translate(md)
        p = doc.paragraphs[0]
        hyperlinks = p._p.findall(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hyperlink"
        )
        assert len(hyperlinks) == 1


# ---------------------------------------------------------------------------
# Lists
# ---------------------------------------------------------------------------

class TestLists:
    def test_ordered_list(self):
        md = "1. first\n2. second\n3. third\n"
        doc = _translate(md)
        items = [
            p for p in doc.paragraphs if p.style.name == "List Number"
        ]
        assert len(items) == 3
        assert items[0].text == "first"

    def test_unordered_list(self):
        md = "- alpha\n- beta\n- gamma\n"
        doc = _translate(md)
        items = [
            p for p in doc.paragraphs if p.style.name == "List Bullet"
        ]
        assert len(items) == 3
        assert items[1].text == "beta"

    def test_list_item_with_bold_lead(self):
        """The Top-N movers format: `- **Name** — body`. Verify the
        bold lead phrase comes through as a bold run inside a bullet
        item."""
        md = "- **Group X** — body of the item\n"
        doc = _translate(md)
        item = next(p for p in doc.paragraphs if p.style.name == "List Bullet")
        bold_runs = [r for r in item.runs if r.bold]
        assert any("Group X" in r.text for r in bold_runs)
        assert "body of the item" in item.text

    def test_separate_ordered_lists_restart_numbering(self):
        """Regression: two ordered lists separated by other content must
        each restart at 1, not share one continuous counter. The bug
        showed up as the "If you read only this page" list opening at 4
        because an earlier three-item list bled its count forward — every
        ordered list inherited the single `numId` carried by the built-in
        "List Number" style, so Word/Google Docs treated them as one
        list."""
        md = (
            "1. first\n2. second\n3. third\n\n"
            "para break\n\n"
            "1. alpha\n2. beta\n"
        )
        doc = _translate(md)
        items = [p for p in doc.paragraphs if p.style.name == "List Number"]
        assert len(items) == 5

        def num_id(p):
            el = p._p.find(".//" + qn("w:numId"))
            return el.get(qn("w:val")) if el is not None else None

        first_list = {num_id(p) for p in items[:3]}
        second_list = {num_id(p) for p in items[3:]}
        # Each list pins its three / two items to a single numId of its own…
        assert len(first_list) == 1 and len(second_list) == 1
        # …and the two lists use *different* numbering instances, which is
        # what makes the second restart at 1.
        assert first_list.isdisjoint(second_list)
        assert next(iter(first_list)) is not None

    def test_ordered_list_honours_start_number(self):
        """A list whose first marker is not 1 (e.g. `3.`) keeps its start
        value via a startOverride rather than being forced back to 1."""
        md = "3. third\n4. fourth\n"
        doc = _translate(md)
        item = next(p for p in doc.paragraphs if p.style.name == "List Number")
        num_id = item._p.find(".//" + qn("w:numId")).get(qn("w:val"))
        numbering = doc.part.numbering_part.element
        num = next(
            n for n in numbering.findall(qn("w:num"))
            if n.get(qn("w:numId")) == num_id
        )
        start = num.find(".//" + qn("w:startOverride"))
        assert start is not None and start.get(qn("w:val")) == "3"


# ---------------------------------------------------------------------------
# Tables
# ---------------------------------------------------------------------------

class TestTables:
    def test_table_structure(self):
        md = (
            "| col1 | col2 | col3 |\n"
            "|---|---|---|\n"
            "| a | b | c |\n"
            "| d | e | f |\n"
        )
        doc = _translate(md)
        assert len(doc.tables) == 1
        t = doc.tables[0]
        assert len(t.rows) == 3
        assert len(t.columns) == 3
        # Header row text
        assert [c.text for c in t.rows[0].cells] == ["col1", "col2", "col3"]
        # Body
        assert [c.text for c in t.rows[1].cells] == ["a", "b", "c"]

    def test_table_header_is_bold(self):
        md = "| h1 | h2 |\n|---|---|\n| body | row |\n"
        doc = _translate(md)
        t = doc.tables[0]
        header_runs = []
        for cell in t.rows[0].cells:
            for p in cell.paragraphs:
                header_runs.extend(p.runs)
        assert all(r.bold for r in header_runs if r.text.strip())

    def test_table_alignment(self):
        """`|---:|` right-aligns; `|:---:|` centres; `|:---|` lefts."""
        md = (
            "| left | right | centre |\n"
            "|:---|---:|:---:|\n"
            "| a | b | c |\n"
        )
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        doc = _translate(md)
        body_row = doc.tables[0].rows[1]
        aligns = [
            body_row.cells[i].paragraphs[0].alignment for i in range(3)
        ]
        # mistune-detected alignments → docx alignment enum values
        assert aligns[0] == WD_ALIGN_PARAGRAPH.LEFT
        assert aligns[1] == WD_ALIGN_PARAGRAPH.RIGHT
        # centre — mistune emits "center" for `|:---:|`
        assert aligns[2] == WD_ALIGN_PARAGRAPH.CENTER


# ---------------------------------------------------------------------------
# Block-level features
# ---------------------------------------------------------------------------

class TestBlockFeatures:
    def test_block_quote_is_italic_and_indented(self):
        md = "> A quoted line of prose.\n"
        doc = _translate(md)
        p = doc.paragraphs[0]
        assert "A quoted line" in p.text
        assert p.paragraph_format.left_indent is not None
        # Force_italic should mean the run is italic
        assert any(r.italic for r in p.runs)

    def test_fenced_code_block_monospace(self):
        md = "```\nline one\nline two\n```\n"
        doc = _translate(md)
        p = doc.paragraphs[0]
        # Content runs (non-whitespace) are Courier New; the inter-line
        # break run is whitespace and doesn't carry font info.
        content_runs = [r for r in p.runs if r.text and r.text.strip()]
        assert all(r.font.name == "Courier New" for r in content_runs)
        assert any("line one" in r.text for r in content_runs)
        assert any("line two" in r.text for r in content_runs)

    def test_thematic_break_emits_paragraph(self):
        md = "Before\n\n---\n\nAfter\n"
        doc = _translate(md)
        # Three paragraphs: "Before", the thematic-break paragraph, "After"
        texts = [p.text for p in doc.paragraphs]
        assert "Before" in texts
        assert "After" in texts


# ---------------------------------------------------------------------------
# Chart injection
# ---------------------------------------------------------------------------

# Minimal valid 1x1 PNG for chart-injection tests — avoids matplotlib
# overhead.
_TINY_PNG = (
    b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00"
    b"\x01\x08\x06\x00\x00\x00\x1f\x15\xc4\x89\x00\x00\x00\rIDATx\x9cc"
    b"\xf8\xff\xff?\x03\x00\x05\xfe\x02\xfe\xa1\x99\x9d\x12\x00\x00\x00"
    b"\x00IEND\xaeB`\x82"
)


class TestChartInjection:
    def test_chart_injects_after_list_item_with_finding_token(self):
        """A list item carrying `finding/N` triggers a picture insert
        when chart_for_finding returns a non-empty list."""
        md = "- **Item** — finding/42 token here\n"
        doc = _translate(
            md,
            chart_for_finding=lambda fid: [_TINY_PNG] if fid == 42 else [],
        )
        assert _images(doc) == 1

    def test_chart_skipped_when_no_finding_token(self):
        md = "- **Item** — no token here\n"
        doc = _translate(
            md, chart_for_finding=lambda fid: [_TINY_PNG],
        )
        assert _images(doc) == 0

    def test_chart_skipped_when_lookup_returns_empty(self):
        md = "- **Item** — finding/42 here\n"
        doc = _translate(md, chart_for_finding=lambda fid: [])
        assert _images(doc) == 0

    def test_chart_skipped_when_lookup_returns_none(self):
        md = "- **Item** — finding/42 here\n"
        doc = _translate(md, chart_for_finding=lambda fid: None)
        assert _images(doc) == 0

    def test_chart_first_occurrence_only(self):
        """The same finding mentioned in two list items → chart set
        inserts only after the first. Real-world: top movers appear
        in both the Top-N section AND in Tier 2 state-of-play."""
        md = (
            "- **First mention** — finding/99 ...\n"
            "- **Second mention** — finding/99 again ...\n"
        )
        doc = _translate(
            md, chart_for_finding=lambda fid: [_TINY_PNG],
        )
        assert _images(doc) == 1

    def test_multiple_charts_per_finding(self):
        """A finding can have multiple charts — lookup returns a list
        and the translator inserts each one after the first occurrence
        of the finding's list item."""
        md = "- **Item** — finding/77 ...\n"
        doc = _translate(
            md,
            chart_for_finding=lambda fid: [_TINY_PNG, _TINY_PNG, _TINY_PNG],
        )
        assert _images(doc) == 3

    def test_chart_not_inserted_when_callable_is_none(self):
        md = "- **Item** — finding/42\n"
        doc = _translate(md)  # no chart_for_finding
        assert _images(doc) == 0


# ---------------------------------------------------------------------------
# Robustness
# ---------------------------------------------------------------------------

class TestRobustness:
    def test_empty_markdown_yields_empty_doc(self):
        doc = _translate("")
        # Document() starts with 0 paragraphs, translator adds none.
        assert len(doc.paragraphs) == 0

    def test_only_blank_lines(self):
        doc = _translate("\n\n\n")
        assert len(doc.paragraphs) == 0

    def test_unknown_inline_falls_back_to_text(self):
        """A markdown construct mistune doesn't recognise should
        degrade to plain text rather than crashing — important for
        forward-compat with future sections."""
        # An auto-link is technically html-tagged by mistune; should
        # come through as text.
        md = "Visit <https://example.com> for info.\n"
        doc = _translate(md)
        assert "example.com" in doc.paragraphs[0].text


# ---------------------------------------------------------------------------
# Metadata-section shading (the blue-tint boilerplate convention)
# ---------------------------------------------------------------------------

from docx.oxml.ns import qn  # noqa: E402


def _is_shaded(p, fill: str = "ABCDEF") -> bool:
    pPr = p._p.pPr
    return pPr is not None and any(
        s.get(qn("w:fill")) == fill for s in pPr.findall(qn("w:shd"))
    )


class TestMetadataShading:
    """Boilerplate/orientation sections are tinted (w:shd fill) so a
    reader of the Doc can tell them from actual findings. A listed
    heading opens shading; a deeper heading stays inside; a same-or-
    higher non-listed heading closes it."""

    MD = (
        "## Orientation\n\nBoilerplate body.\n\n"
        "### Sub-orientation\n\nStill boilerplate.\n\n"
        "## Real content\n\nA finding.\n"
    )

    def test_listed_section_and_subsections_are_shaded(self):
        doc = _translate(
            self.MD,
            shaded_section_headings={"Orientation"},
            shade_fill="ABCDEF",
        )
        by_text = {p.text: _is_shaded(p) for p in doc.paragraphs if p.text}
        assert by_text["Orientation"] is True
        assert by_text["Boilerplate body."] is True
        assert by_text["Sub-orientation"] is True
        assert by_text["Still boilerplate."] is True

    def test_same_level_unlisted_heading_closes_shading(self):
        doc = _translate(
            self.MD,
            shaded_section_headings={"Orientation"},
            shade_fill="ABCDEF",
        )
        by_text = {p.text: _is_shaded(p) for p in doc.paragraphs if p.text}
        assert by_text["Real content"] is False
        assert by_text["A finding."] is False

    def test_reading_the_numbers_registered_in_both_doc_sets(self):
        """The 'Reading the numbers' key (briefing_pack._helpers.
        _reading_the_numbers_md) is orientation material — it must carry
        the boilerplate tint in BOTH the findings and the leads Doc.
        Guards against the key silently losing its tint (and, because a
        same-level unlisted heading closes shading, against it cutting
        the tint off the 'Predictability badges' block that follows it
        in the findings doc)."""
        from briefing_pack.docx import (
            _LEADS_METADATA_HEADINGS,
            _METADATA_SECTION_HEADINGS,
        )
        assert "Reading the numbers" in _METADATA_SECTION_HEADINGS
        assert "Reading the numbers" in _LEADS_METADATA_HEADINGS
