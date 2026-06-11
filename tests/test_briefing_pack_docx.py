"""Tests for the docx output module (`briefing_pack.docx`).

Pure-function tests live first (no DB), followed by integration tests
that seed the test DB and run the full renderer end-to-end.

Determinism testing notes (design doc step 5):
- docx-internal timestamps and openpyxl metadata are NOT byte-stable
  across runs. Don't SHA256-hash the entire .docx — the test will be
  flaky.
- matplotlib Agg with explicit color overrides and explicit DPI IS
  byte-stable on a given host. PNG-bytes tests are safe there.
- Cross-host PNG byte stability is not guaranteed (font fallback,
  freetype version drift). For CI we'd need to pin a font; out of
  scope for v1.
"""

from __future__ import annotations

import hashlib
import json
from datetime import date
from pathlib import Path

import psycopg2
import pytest

from docx import Document

import briefing_pack
import briefing_pack.docx as bp_docx
from tests.test_briefing_pack import (
    _seed_eurostat_release,
    _seed_gacc_bilateral_aggregate_yoy_finding,
    _seed_hs_yoy_finding,
    _seed_run,
)


@pytest.fixture(autouse=True)
def _direct_db_url(test_db_url, monkeypatch):
    monkeypatch.setenv("DATABASE_URL", test_db_url)
    monkeypatch.delenv(briefing_pack.PERMALINK_BASE_ENV, raising=False)


@pytest.fixture
def empty_findings(test_db_url):
    with psycopg2.connect(test_db_url) as conn, conn.cursor() as cur:
        cur.execute(
            "TRUNCATE findings, observations, source_snapshots, "
            "eurostat_raw_rows, scrape_runs, releases "
            "RESTART IDENTITY CASCADE"
        )
    yield


# ---------------------------------------------------------------------------
# Pure-function tests
# ---------------------------------------------------------------------------

class TestPickEurScale:
    def test_billions(self):
        scale, label = bp_docx._pick_eur_scale(2.5e9)
        assert scale == 1e9
        assert label == "€ billions"

    def test_millions(self):
        scale, label = bp_docx._pick_eur_scale(450e6)
        assert scale == 1e6
        assert label == "€ millions"

    def test_thousands(self):
        scale, label = bp_docx._pick_eur_scale(5e3)
        assert scale == 1e3
        assert label == "€ thousands"

    def test_units(self):
        scale, label = bp_docx._pick_eur_scale(500)
        assert scale == 1.0
        assert label == "€"

    def test_billions_boundary_exactly_1e9(self):
        # 1e9 should be billions (≥, not >)
        scale, _ = bp_docx._pick_eur_scale(1e9)
        assert scale == 1e9


class TestMonthsBack:
    def test_23_months_back_from_feb_2026(self):
        # Real-world case: current_end = 2026-02-01, 24-month window
        # starts at March 2024 (23 months earlier).
        assert bp_docx._months_back(date(2026, 2, 1), 23) == date(2024, 3, 1)

    def test_1_month_back_across_year_boundary(self):
        assert bp_docx._months_back(date(2026, 1, 1), 1) == date(2025, 12, 1)

    def test_12_months_back_lands_on_same_month_prior_year(self):
        assert bp_docx._months_back(date(2026, 5, 1), 12) == date(2025, 5, 1)

    def test_zero_months_back_is_identity(self):
        assert bp_docx._months_back(date(2026, 5, 1), 0) == date(2026, 5, 1)


class TestMonthIter:
    def test_yields_12_months_inclusive(self):
        months = list(bp_docx._month_iter(date(2025, 3, 1), date(2026, 2, 1)))
        assert len(months) == 12
        assert months[0] == date(2025, 3, 1)
        assert months[-1] == date(2026, 2, 1)

    def test_single_month(self):
        months = list(bp_docx._month_iter(date(2026, 5, 1), date(2026, 5, 1)))
        assert months == [date(2026, 5, 1)]

    def test_spans_year_boundary(self):
        months = list(bp_docx._month_iter(date(2025, 11, 1), date(2026, 2, 1)))
        assert months == [
            date(2025, 11, 1), date(2025, 12, 1),
            date(2026, 1, 1), date(2026, 2, 1),
        ]


class TestFlowLabel:
    def test_export_subkind(self):
        assert bp_docx._flow_label_for_subkind("hs_group_yoy_export") == \
            "EU-27 exports to China"

    def test_import_subkind(self):
        assert bp_docx._flow_label_for_subkind("hs_group_yoy") == \
            "EU-27 imports from China"


class TestBuildBilateralSummaryBarPng:
    """Two-bar prior-vs-current chart for gacc_bilateral_aggregate_yoy*
    findings. Editorial purpose: turn the diff section's percentages
    into € magnitudes at a glance."""

    def test_returns_png_bytes(self):
        png = bp_docx._build_bilateral_summary_bar_png(
            partner_label="Germany",
            subkind="gacc_bilateral_aggregate_yoy",
            current_end=date(2026, 2, 1),
            current_eur=8.34e9,
            prior_eur=14.07e9,
            yoy_pct=-0.407,
        )
        assert png.startswith(b"\x89PNG\r\n\x1a\n")
        assert len(png) > 5_000

    def test_handles_positive_yoy(self):
        png = bp_docx._build_bilateral_summary_bar_png(
            partner_label="ASEAN",
            subkind="gacc_bilateral_aggregate_yoy_import",
            current_end=date(2026, 2, 1),
            current_eur=12.0e9, prior_eur=10.0e9, yoy_pct=0.20,
        )
        assert png.startswith(b"\x89PNG\r\n\x1a\n")

    def test_handles_zero_prior(self):
        """Edge case — a partner with no prior period flow (e.g. a
        new aggregate). Chart still produces, prior bar at 0."""
        png = bp_docx._build_bilateral_summary_bar_png(
            partner_label="X",
            subkind="gacc_bilateral_aggregate_yoy",
            current_end=date(2026, 2, 1),
            current_eur=1e9, prior_eur=0.0, yoy_pct=999.0,
        )
        assert png is not None

    def test_deterministic(self):
        kwargs = dict(
            partner_label="Germany",
            subkind="gacc_bilateral_aggregate_yoy",
            current_end=date(2026, 2, 1),
            current_eur=8.34e9, prior_eur=14.07e9, yoy_pct=-0.407,
        )
        a = bp_docx._build_bilateral_summary_bar_png(**kwargs)
        b = bp_docx._build_bilateral_summary_bar_png(**kwargs)
        assert hashlib.sha256(a).hexdigest() == hashlib.sha256(b).hexdigest()


class TestComputeTopBilateralMovers:
    """Integration tests against the test DB for the
    _compute_top_bilateral_movers selector."""

    def test_empty_db_returns_empty(self, empty_findings, test_db_url):
        with psycopg2.connect(test_db_url) as conn:
            cur = conn.cursor(cursor_factory=psycopg2.extras.DictCursor)
            assert bp_docx._compute_top_bilateral_movers(cur, limit=5) == []

    def test_respects_5pp_threshold(self, empty_findings, test_db_url):
        """Findings with |yoy_pct| < 5% must be excluded — matches
        Tier 1 diff's material-shift cutoff."""
        with psycopg2.connect(test_db_url) as conn:
            cur = conn.cursor()
            run = _seed_run(cur)
            _seed_gacc_bilateral_aggregate_yoy_finding(
                cur, run, "Germany", yoy_pct=0.30,
                period=date(2026, 2, 1),
            )
            _seed_gacc_bilateral_aggregate_yoy_finding(
                cur, run, "France", yoy_pct=0.02,  # below threshold
                period=date(2026, 2, 1),
            )
            conn.commit()
        with psycopg2.connect(test_db_url) as conn:
            cur = conn.cursor(cursor_factory=psycopg2.extras.DictCursor)
            results = bp_docx._compute_top_bilateral_movers(cur, limit=5)
        partners = [r["partner_label"] for r in results]
        assert "Germany" in partners
        assert "France" not in partners

    def test_ranks_by_absolute_yoy(self, empty_findings, test_db_url):
        """A -40% move ranks above a +30% move; absolute value matters,
        not sign."""
        with psycopg2.connect(test_db_url) as conn:
            cur = conn.cursor()
            run = _seed_run(cur)
            _seed_gacc_bilateral_aggregate_yoy_finding(
                cur, run, "DropPartner", yoy_pct=-0.40,
                period=date(2026, 2, 1),
            )
            _seed_gacc_bilateral_aggregate_yoy_finding(
                cur, run, "RisePartner", yoy_pct=0.30,
                period=date(2026, 2, 1),
            )
            conn.commit()
        with psycopg2.connect(test_db_url) as conn:
            cur = conn.cursor(cursor_factory=psycopg2.extras.DictCursor)
            results = bp_docx._compute_top_bilateral_movers(cur, limit=5)
        assert results[0]["partner_label"] == "DropPartner"
        assert results[1]["partner_label"] == "RisePartner"

    def test_only_latest_anchor(self, empty_findings, test_db_url):
        """Older-anchor findings excluded — only the latest cycle's
        bilaterals are charted."""
        with psycopg2.connect(test_db_url) as conn:
            cur = conn.cursor()
            run = _seed_run(cur)
            _seed_gacc_bilateral_aggregate_yoy_finding(
                cur, run, "OldPartner", yoy_pct=0.50,
                period=date(2025, 8, 1),  # older
            )
            _seed_gacc_bilateral_aggregate_yoy_finding(
                cur, run, "NewPartner", yoy_pct=0.10,
                period=date(2026, 2, 1),  # latest
            )
            conn.commit()
        with psycopg2.connect(test_db_url) as conn:
            cur = conn.cursor(cursor_factory=psycopg2.extras.DictCursor)
            results = bp_docx._compute_top_bilateral_movers(cur, limit=5)
        partners = [r["partner_label"] for r in results]
        assert "NewPartner" in partners
        assert "OldPartner" not in partners

    def test_respects_limit(self, empty_findings, test_db_url):
        with psycopg2.connect(test_db_url) as conn:
            cur = conn.cursor()
            run = _seed_run(cur)
            for i, p in enumerate(["A", "B", "C", "D", "E"]):
                _seed_gacc_bilateral_aggregate_yoy_finding(
                    cur, run, p, yoy_pct=0.10 + i * 0.05,
                    period=date(2026, 2, 1),
                )
            conn.commit()
        with psycopg2.connect(test_db_url) as conn:
            cur = conn.cursor(cursor_factory=psycopg2.extras.DictCursor)
            results = bp_docx._compute_top_bilateral_movers(cur, limit=3)
        assert len(results) == 3


def test_render_findings_docx_injects_chart_for_bilateral_finding(
    empty_findings, test_db_url, tmp_path,
):
    """A material-shift bilateral finding seeded → its
    `finding/N` token in the markdown's Tier 1 diff triggers chart
    injection. We verify end-to-end that bilaterals get charts
    alongside top-mover hs_group_yoy findings."""
    out_path = tmp_path / "bilateral.docx"
    with psycopg2.connect(test_db_url) as conn:
        cur = conn.cursor()
        run = _seed_run(cur)
        _seed_eurostat_release(cur, date(2026, 2, 1))
        # Seed a bilateral with a strong YoY so it passes the 5pp
        # filter + ranks high.
        _seed_gacc_bilateral_aggregate_yoy_finding(
            cur, run, "Germany", yoy_pct=-0.31, period=date(2026, 2, 1),
        )
        # Also need at least one current-cycle eurostat release for
        # render() to produce a sensible markdown shape.
        conn.commit()

    bp_docx.render_findings_docx(out_path)
    doc = Document(str(out_path))

    # At least one chart for the seeded bilateral finding
    assert _count_images(doc) >= 1


class TestBuildPerReporterBarPng:
    """Per-reporter grouped-bar chart for top movers — answers
    'which country is driving the move?' alongside the trajectory
    line chart."""

    def _breakdown(self) -> list[dict]:
        return [
            {"reporter": "Germany", "current_eur": 4.21e9, "prior_eur": 2.48e9,
             "yoy_pct": -0.41, "share_of_group_delta_pct": 0.42},
            {"reporter": "France", "current_eur": 1.92e9, "prior_eur": 1.13e9,
             "yoy_pct": -0.41, "share_of_group_delta_pct": 0.18},
            {"reporter": "Italy", "current_eur": 0.85e9, "prior_eur": 0.51e9,
             "yoy_pct": -0.40, "share_of_group_delta_pct": 0.08},
        ]

    def test_returns_png_bytes(self):
        png = bp_docx._build_per_reporter_bar_png(
            breakdown=self._breakdown(),
            group_name="Finished cars (broad)",
            flow_label="EU-27 exports to China",
        )
        assert png is not None
        assert png.startswith(b"\x89PNG\r\n\x1a\n")
        assert len(png) > 5_000

    def test_returns_none_on_empty_breakdown(self):
        assert bp_docx._build_per_reporter_bar_png(
            breakdown=[],
            group_name="X",
            flow_label="Y",
        ) is None

    def test_returns_none_when_all_rows_have_no_values(self):
        assert bp_docx._build_per_reporter_bar_png(
            breakdown=[
                {"reporter": "DE", "current_eur": None, "prior_eur": None},
                {"reporter": "FR", "current_eur": None, "prior_eur": None},
            ],
            group_name="X",
            flow_label="Y",
        ) is None

    def test_top_k_cap(self):
        """Even with more reporters available, the chart selects top-K
        by absolute delta. We can't introspect the chart pixels from
        a test, but we can call with K=2 and verify the chart still
        produces (i.e., doesn't error on slicing)."""
        breakdown = self._breakdown() * 3  # 9 rows
        png = bp_docx._build_per_reporter_bar_png(
            breakdown=breakdown,
            group_name="X",
            flow_label="Y",
            top_k=2,
        )
        assert png is not None
        assert png.startswith(b"\x89PNG\r\n\x1a\n")

    def test_deterministic(self):
        kwargs = dict(
            breakdown=self._breakdown(),
            group_name="Finished cars",
            flow_label="EU-27 exports to China",
        )
        a = bp_docx._build_per_reporter_bar_png(**kwargs)
        b = bp_docx._build_per_reporter_bar_png(**kwargs)
        assert hashlib.sha256(a).hexdigest() == hashlib.sha256(b).hexdigest()


class TestBuildChartPng:
    def _series(self) -> dict[date, float]:
        return {
            date(2024, 3 + i, 1) if (3 + i) <= 12 else date(2024 + (3 + i - 1) // 12, (3 + i - 1) % 12 + 1, 1): float(100e6 + i * 5e6)
            for i in range(24)
        }

    def test_returns_png_bytes(self):
        series = self._series()
        png = bp_docx._build_chart_png(
            current_end=date(2026, 2, 1),
            monthly_eur=series,
            group_name="Test group",
            flow_label="EU-27 imports from China",
        )
        assert isinstance(png, bytes)
        # PNG magic bytes
        assert png.startswith(b"\x89PNG\r\n\x1a\n")
        # Reasonable size — not a one-pixel placeholder
        assert len(png) > 5_000

    def test_deterministic_same_input_same_output(self):
        """Same series → same PNG bytes. matplotlib Agg with our color +
        DPI pins is deterministic on a given host. This test catches
        regressions where a future contributor adds non-determinism
        (e.g. random palette, default ax title with timestamp)."""
        series = self._series()
        kwargs = dict(
            current_end=date(2026, 2, 1),
            monthly_eur=series,
            group_name="Test group",
            flow_label="EU-27 imports from China",
        )
        a = bp_docx._build_chart_png(**kwargs)
        b = bp_docx._build_chart_png(**kwargs)
        assert hashlib.sha256(a).hexdigest() == hashlib.sha256(b).hexdigest()

    def test_handles_empty_series(self):
        """Empty data dict — chart still renders (no exception) and
        produces a valid PNG showing an empty plot. Real-world case:
        a finding whose eurostat_raw_rows are missing in the test DB."""
        png = bp_docx._build_chart_png(
            current_end=date(2026, 2, 1),
            monthly_eur={},
            group_name="No data",
            flow_label="EU-27 imports from China",
        )
        assert png.startswith(b"\x89PNG\r\n\x1a\n")

    def test_handles_gaps_in_series(self):
        """Some months missing — matplotlib NaN handling skips them
        rather than raising. Important for real findings where Eurostat
        partial-window caveats apply."""
        series = {
            date(2024, 3, 1): 100e6,
            date(2024, 5, 1): 110e6,  # April missing
            date(2024, 9, 1): 120e6,  # May–Aug missing
        }
        png = bp_docx._build_chart_png(
            current_end=date(2026, 2, 1),
            monthly_eur=series,
            group_name="Sparse data",
            flow_label="EU-27 imports from China",
        )
        assert png.startswith(b"\x89PNG\r\n\x1a\n")


# ---------------------------------------------------------------------------
# Integration tests — full renderer against a seeded DB
# ---------------------------------------------------------------------------

def _seed_chart_capable_finding(
    cur,
    run_id: int,
    group_name: str,
    *,
    subkind: str = "hs_group_yoy",
    hs_pattern: str = "8507%",
    **kwargs,
) -> int:
    """Wrap `_seed_hs_yoy_finding` and patch the detail JSONB to include
    the `method_query.flow` / `method_query.partners` /
    `method_query.hs_patterns` fields the chart fetcher needs.

    The shared `_seed_hs_yoy_finding` helper omits these because the
    older test surface (markdown rendering) doesn't need them; the
    docx renderer's chart-fetch path does. Production findings always
    have all three (set by `anomalies.detect_hs_group_yoy`)."""
    fid = _seed_hs_yoy_finding(
        cur, run_id, group_name, subkind=subkind, **kwargs,
    )
    cur.execute("SELECT detail FROM findings WHERE id = %s", (fid,))
    detail = cur.fetchone()[0]
    detail.setdefault("method_query", {})
    detail["method_query"]["flow"] = (
        2 if subkind.endswith("_export") else 1
    )
    detail["method_query"]["partners"] = ["CN", "HK", "MO"]
    detail["method_query"]["hs_patterns"] = [hs_pattern]
    cur.execute(
        "UPDATE findings SET detail = %s::jsonb WHERE id = %s",
        (json.dumps(detail), fid),
    )
    return fid


def _seed_eurostat_raw_rows_for_finding(
    cur,
    run_id: int,
    *,
    hs_pattern: str,
    reporter: str = "DE",
    partner: str = "CN",
    flow: int = 1,
    anchor: date = date(2026, 2, 1),
    months: int = 24,
    monthly_eur: float = 100e6,
) -> None:
    """Seed `months` of eurostat_raw_rows ending at `anchor` so the chart
    fetch in `_fetch_monthly_eur_series` returns non-empty data."""
    code = hs_pattern.rstrip("%")
    for i in range(months):
        period = bp_docx._months_back(anchor, months - 1 - i)
        cur.execute(
            """
            INSERT INTO eurostat_raw_rows
                (scrape_run_id, period, reporter, partner, product_nc,
                 flow, value_eur, quantity_kg)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
            """,
            (run_id, period, reporter, partner, code, flow,
             monthly_eur * (1.0 + 0.02 * i),  # gentle upward trend
             monthly_eur * 0.5),
        )


def _count_images(doc: Document) -> int:
    return sum(
        1 for p in doc.paragraphs
        for r in p.runs
        if r.element.xpath(".//w:drawing")
    )


def test_render_findings_docx_smoke(empty_findings, test_db_url, tmp_path):
    """Seed one eligible mover + its raw rows, render, and assert the
    docx file is produced and opens as a valid Document with the full
    findings.md content rendered into it."""
    out_path = tmp_path / "smoke.docx"
    with psycopg2.connect(test_db_url) as conn:
        cur = conn.cursor()
        run = _seed_run(cur)
        _seed_eurostat_release(cur, date(2026, 2, 1))
        _seed_chart_capable_finding(
            cur, run, "EV batteries (Li-ion)",
            yoy_pct=0.35, current_eur=27e9, prior_eur=20e9, low_base=False,
        )
        _seed_eurostat_raw_rows_for_finding(cur, run, hs_pattern="8507%")
        conn.commit()

    result = bp_docx.render_findings_docx(out_path)
    assert result == out_path
    assert out_path.exists()

    doc = Document(str(out_path))
    # v4: full markdown content → many paragraphs (period coverage,
    # scope notes, findings inventory, top movers, methodology footer,
    # etc.) plus at least one chart for the seeded mover.
    assert len(doc.paragraphs) > 20


def test_render_findings_docx_full_content_parity(
    empty_findings, test_db_url, tmp_path,
):
    """v4 contract: the docx contains the same content as findings.md.
    Verify by checking that several known sections from the markdown's
    structural skeleton appear as headings in the docx."""
    out_path = tmp_path / "parity.docx"
    with psycopg2.connect(test_db_url) as conn:
        cur = conn.cursor()
        run = _seed_run(cur)
        _seed_eurostat_release(cur, date(2026, 2, 1))
        _seed_chart_capable_finding(
            cur, run, "EV batteries (Li-ion)",
            yoy_pct=0.345, current_eur=27.25e9, prior_eur=20.27e9,
        )
        _seed_eurostat_raw_rows_for_finding(cur, run, hs_pattern="8507%")
        conn.commit()

    bp_docx.render_findings_docx(out_path)
    doc = Document(str(out_path))

    headings = [
        p.text.strip()
        for p in doc.paragraphs
        if p.style and p.style.name.startswith("Heading")
    ]
    # The renderer always emits these structural sections regardless
    # of content. Catches regressions where the translator stops
    # walking blocks midway.
    assert any("Scope notes" in h for h in headings), \
        f"expected 'Scope notes' heading, got: {headings[:20]}"
    assert any("Findings included" in h for h in headings), \
        f"expected 'Findings included' heading, got: {headings[:20]}"
    assert any("Methodology" in h for h in headings), \
        f"expected a Methodology heading, got: {headings[:20]}"


def test_render_findings_docx_injects_chart_for_top_mover(
    empty_findings, test_db_url, tmp_path,
):
    """A seeded top mover with underlying raw rows → exactly one chart
    image in the docx (the one in the Top-N section; first-occurrence-
    only guard prevents duplicates in Tier 2 state-of-play)."""
    out_path = tmp_path / "chart.docx"
    with psycopg2.connect(test_db_url) as conn:
        cur = conn.cursor()
        run = _seed_run(cur)
        _seed_eurostat_release(cur, date(2026, 2, 1))
        _seed_chart_capable_finding(
            cur, run, "EV batteries (Li-ion)",
            yoy_pct=0.345, current_eur=27.25e9, prior_eur=20.27e9,
        )
        _seed_eurostat_raw_rows_for_finding(cur, run, hs_pattern="8507%")
        conn.commit()

    bp_docx.render_findings_docx(out_path)
    doc = Document(str(out_path))

    assert _count_images(doc) == 1


def test_render_findings_docx_injects_per_reporter_bar_chart(
    empty_findings, test_db_url, tmp_path,
):
    """A top mover whose detail carries `per_reporter_breakdown` gets
    a second chart (the per-reporter grouped bar) alongside the
    trajectory line chart — answers 'which country is driving the
    move?' alongside the headline trajectory."""
    out_path = tmp_path / "two-charts.docx"
    with psycopg2.connect(test_db_url) as conn:
        cur = conn.cursor()
        run = _seed_run(cur)
        _seed_eurostat_release(cur, date(2026, 2, 1))
        _seed_chart_capable_finding(
            cur, run, "EV batteries (Li-ion)",
            yoy_pct=0.345, current_eur=27.25e9, prior_eur=20.27e9,
            per_reporter_breakdown=[
                {"reporter": "Germany", "current_eur": 12e9, "prior_eur": 9e9,
                 "yoy_pct": 0.33, "share_of_group_delta_pct": 0.43},
                {"reporter": "France", "current_eur": 5e9, "prior_eur": 4e9,
                 "yoy_pct": 0.25, "share_of_group_delta_pct": 0.14},
                {"reporter": "Italy", "current_eur": 3e9, "prior_eur": 2.5e9,
                 "yoy_pct": 0.20, "share_of_group_delta_pct": 0.07},
            ],
        )
        _seed_eurostat_raw_rows_for_finding(cur, run, hs_pattern="8507%")
        conn.commit()

    bp_docx.render_findings_docx(out_path)
    doc = Document(str(out_path))

    # Two charts for this single finding: trajectory line + reporter bar
    assert _count_images(doc) == 2


def test_render_findings_docx_skips_chart_for_finding_without_raw_rows(
    empty_findings, test_db_url, tmp_path,
):
    """A finding seeded without underlying eurostat_raw_rows → no
    chart injected (chart-data fetch returns None; translator's
    chart_for_finding lookup returns None; no picture added).
    The rest of the markdown still renders normally."""
    out_path = tmp_path / "no-chart.docx"
    with psycopg2.connect(test_db_url) as conn:
        cur = conn.cursor()
        run = _seed_run(cur)
        _seed_eurostat_release(cur, date(2026, 2, 1))
        _seed_hs_yoy_finding(
            cur, run, "EV batteries (Li-ion)",
            yoy_pct=0.35, current_eur=27e9, prior_eur=20e9, low_base=False,
        )
        # Deliberately no raw rows + no chart_capable_finding patching,
        # so method_query lacks flow/partners — chart fetch returns None.
        conn.commit()

    bp_docx.render_findings_docx(out_path)
    doc = Document(str(out_path))

    assert _count_images(doc) == 0
    # But the document still has substantive content
    assert len(doc.paragraphs) > 10


def test_render_findings_docx_respects_top_n_for_charts(
    empty_findings, test_db_url, tmp_path,
):
    """`top_n=2` caps the number of charts (not headings — headings
    come from the markdown's structural skeleton). The .md still
    renders all eligible movers in its Top-N section; only the chart
    insertion is bounded by `top_n`."""
    out_path = tmp_path / "topn.docx"
    with psycopg2.connect(test_db_url) as conn:
        cur = conn.cursor()
        run = _seed_run(cur)
        _seed_eurostat_release(cur, date(2026, 2, 1))
        _seed_chart_capable_finding(
            cur, run, "EV batteries (Li-ion)",
            yoy_pct=0.35, current_eur=27e9, prior_eur=20e9,
        )
        _seed_chart_capable_finding(
            cur, run, "Drones and unmanned aircraft",
            yoy_pct=0.40, current_eur=1.0e9, prior_eur=0.7e9,
        )
        _seed_chart_capable_finding(
            cur, run, "Wind generating sets only",
            yoy_pct=0.345, current_eur=375e6, prior_eur=279e6,
        )
        _seed_eurostat_raw_rows_for_finding(cur, run, hs_pattern="8507%")
        conn.commit()

    bp_docx.render_findings_docx(out_path, top_n=2)
    doc = Document(str(out_path))

    assert _count_images(doc) == 2


def test_render_findings_docx_renders_with_no_findings(
    empty_findings, test_db_url, tmp_path,
):
    """Empty DB — markdown render still emits its framing sections
    (period coverage, scope notes, etc.) so the docx is non-empty
    even when no top movers exist. No charts in this case."""
    out_path = tmp_path / "empty.docx"
    # empty_findings has truncated tables; no seeds.

    bp_docx.render_findings_docx(out_path)
    doc = Document(str(out_path))

    # Framing content present (markdown always emits these)
    assert len(doc.paragraphs) > 5
    # No charts because no top movers
    assert _count_images(doc) == 0


def test_render_findings_docx_applies_page_setup(
    empty_findings, test_db_url, tmp_path,
):
    """A4 + 10mm margins must reach the saved .docx — these were the
    fix that prompted the spike, and a regression here would put us
    back to US Letter / 1-inch margins for Lisa."""
    out_path = tmp_path / "pagesetup.docx"
    bp_docx.render_findings_docx(out_path)
    doc = Document(str(out_path))

    section = doc.sections[0]
    # python-docx stores Length objects with `.mm` property
    assert round(section.page_width.mm) == 210
    assert round(section.page_height.mm) == 297
    assert round(section.top_margin.mm) == 10
    assert round(section.bottom_margin.mm) == 10
    assert round(section.left_margin.mm) == 10
    assert round(section.right_margin.mm) == 10
