"""
Drive fidelity spike — local-only generator (no OAuth, no upload).

Generates a sample .docx and .xlsx that mirror the shape of the real
briefing-pack bundle (Top 5 movers, per-finding cards, native chart
in the spreadsheet) so we can eyeball Drive → Docs / Sheets
conversion fidelity by uploading manually.

Usage:
    python scripts/drive_spike_local.py

Output:
    exports/spike-2026-05-16/03_Findings_test.docx
    exports/spike-2026-05-16/04_Data_test.xlsx

What this is testing (legs 1 and 2 of the spike spec):
- Does python-docx → upload to Drive → open with Google Docs preserve
  headings, paragraphs, tables, bullet lists, inline images, and
  enough emoji/Unicode for the predictability badges (🟡 🔴 🟢 etc.)?
- Does openpyxl-authored .xlsx with a native LineChart → upload to
  Drive → open with Sheets preserve the chart as a real editable
  Sheets chart object (not a flat image)?

Content lifted/adapted from exports/2026-05-15-1811/03_Findings.md for
realism. Time-series values are synthesised — they look like a real
shift but are not pulled from the database. This is a fidelity test,
not a content test.
"""

from __future__ import annotations

import io
from pathlib import Path

import matplotlib

matplotlib.use("Agg")  # non-interactive backend
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Mm, Pt
from openpyxl import Workbook
from openpyxl.chart import BarChart, LineChart, Reference
from openpyxl.styles import Font

OUT_DIR = Path(__file__).parent.parent / "exports" / "spike-2026-05-16"


# ---------------------------------------------------------------------------
# Synthesised time-series data — shaped to look like the real findings
# in exports/2026-05-15-1811/03_Findings.md but not pulled from the DB.
# ---------------------------------------------------------------------------

MONTHS = [
    "Mar 25", "Apr 25", "May 25", "Jun 25", "Jul 25", "Aug 25",
    "Sep 25", "Oct 25", "Nov 25", "Dec 25", "Jan 26", "Feb 26",
]

# EV batteries (Li-ion), EU-27 imports CN→reporter. Real finding said
# +34.5% YoY to €27.25B 12mo. Made-up monthly series with rising
# current vs flat prior.
EV_BATTERY_PRIOR = [
    1.55, 1.60, 1.62, 1.58, 1.65, 1.71, 1.68, 1.74, 1.79, 1.82, 1.61, 1.70,
]
EV_BATTERY_CURRENT = [
    1.95, 2.10, 2.18, 2.20, 2.31, 2.42, 2.38, 2.51, 2.62, 2.70, 2.20, 2.28,
]

# Finished cars (broad), EU-27 exports reporter→CN. Real finding said
# -40.7% YoY to €8.34B 12mo. Falling current vs higher prior.
FINISHED_CARS_PRIOR = [
    1.20, 1.15, 1.18, 1.12, 1.08, 1.05, 1.00, 0.98, 0.95, 0.92, 0.88, 0.82,
]
FINISHED_CARS_CURRENT = [
    0.80, 0.75, 0.72, 0.68, 0.65, 0.62, 0.58, 0.60, 0.55, 0.52, 0.50, 0.48,
]


# ---------------------------------------------------------------------------
# Chart generation — returns PNG bytes for embedding in docx.
# ---------------------------------------------------------------------------

def line_chart_png(
    months: list[str],
    prior: list[float],
    current: list[float],
    *,
    title: str,
    ylabel: str,
) -> bytes:
    fig, ax = plt.subplots(figsize=(6.5, 3.2), dpi=150)
    ax.plot(months, prior, label="Prior 12mo", linewidth=2, color="#888888")
    ax.plot(months, current, label="Current 12mo", linewidth=2.5, color="#cc3333")
    ax.set_title(title, fontsize=11, loc="left")
    ax.set_ylabel(ylabel, fontsize=9)
    ax.tick_params(axis="x", labelsize=7, rotation=45)
    ax.tick_params(axis="y", labelsize=8)
    ax.grid(True, alpha=0.3)
    ax.legend(loc="best", fontsize=8, frameon=False)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    plt.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format="png", bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf.read()


def grouped_bar_png(
    labels: list[str],
    series_a: list[float],
    series_b: list[float],
    *,
    series_a_name: str,
    series_b_name: str,
    title: str,
    ylabel: str,
) -> bytes:
    fig, ax = plt.subplots(figsize=(6.5, 3.2), dpi=150)
    x = range(len(labels))
    width = 0.35
    ax.bar([i - width / 2 for i in x], series_a, width, label=series_a_name, color="#888888")
    ax.bar([i + width / 2 for i in x], series_b, width, label=series_b_name, color="#cc3333")
    ax.set_xticks(list(x))
    ax.set_xticklabels(labels, fontsize=9)
    ax.set_title(title, fontsize=11, loc="left")
    ax.set_ylabel(ylabel, fontsize=9)
    ax.tick_params(axis="y", labelsize=8)
    ax.grid(True, axis="y", alpha=0.3)
    ax.legend(loc="best", fontsize=8, frameon=False)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    plt.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format="png", bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf.read()


# ---------------------------------------------------------------------------
# Docx generation
# ---------------------------------------------------------------------------

def build_docx(out_path: Path) -> None:
    doc = Document()

    # Page setup: A4 portrait, 10mm margins all sides. python-docx applies
    # these to the document's default section. Note that Google Docs' own
    # "Pageless" mode is a Docs-side toggle that isn't expressible in .docx —
    # set it after conversion via File → Page setup → Pageless if wanted.
    section = doc.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.top_margin = Mm(10)
    section.bottom_margin = Mm(10)
    section.left_margin = Mm(10)
    section.right_margin = Mm(10)

    # Set body font size baseline so headings inherit a sensible scale
    style = doc.styles["Normal"]
    style.font.size = Pt(11)

    doc.add_heading("Meridian — Findings (spike test export)", level=0)

    p = doc.add_paragraph()
    p.add_run("Period: ").bold = True
    p.add_run("12-month window ending 2026-02-01. Content lifted from a real cycle for realism; ")
    p.add_run("time-series values are synthesised — fidelity test, not data test.").italic = True

    # ---- Top movers section ----
    doc.add_heading("Top 5 movers this cycle", level=1)

    italic_preamble = doc.add_paragraph()
    italic_preamble.add_run(
        "Editorially-quotable shifts ranked by a composite of |YoY| × log(€). "
        "Filters: ≥10pp move, ≥€100M current 12mo total, not low-base, predictability "
        "badge ≠ 🔴 (no badge is fine — groups without enough T-6 history yet are still "
        "eligible). Drill into each via its Tier 2 anchor."
    ).italic = True

    movers = [
        ("Finished cars (broad) 🟡", "EU-27 exports (reporter→CN): -40.7% (kg -38.4%) to €8.34B"),
        ("EV batteries (Li-ion) 🟡", "EU-27 imports (CN→reporter): +34.5% (kg +69.4%) to €27.25B"),
        ("Drones and unmanned aircraft 🟡", "EU-27 imports (CN→reporter): +39.2% (kg +55.7%) to €1.10B"),
        ("EV batteries (Li-ion) 🟡", "EU-27 exports (reporter→CN): -36.2% (kg -8.7%) to €453.7M"),
        ("Wind generating sets only", "EU-27 imports (CN→reporter): +34.5% (kg +55.5%) to €375.1M"),
    ]
    for i, (name, desc) in enumerate(movers, start=1):
        p = doc.add_paragraph(style="List Number")
        p.add_run(f"{name} — ").bold = True
        p.add_run(desc)

    # ---- Featured finding 1: line chart ----
    doc.add_heading("EV batteries (Li-ion) — EU-27 imports from China", level=1)

    doc.add_paragraph(
        "China's exports of Li-ion EV batteries to the EU-27 reached €27.25B "
        "over the 12 months to February 2026 — up 34.5% in value terms and "
        "69.4% in kilogrammes year-on-year. The volume-vs-value gap is "
        "consistent with the falling per-unit price trend Soapbox Trade has "
        "documented across H2 2025."
    )

    chart_png = line_chart_png(
        MONTHS,
        EV_BATTERY_PRIOR,
        EV_BATTERY_CURRENT,
        title="EU-27 imports of Li-ion EV batteries from China (€ billions, monthly)",
        ylabel="€ billions",
    )
    doc.add_picture(io.BytesIO(chart_png), width=Mm(190))

    # ---- A representative table ----
    doc.add_heading("Per-window breakdown", level=2)

    doc.add_paragraph(
        "Latest four 12-month windows, showing how the YoY has tightened "
        "since the Jan-Feb combined-release fix shipped (2026-05-15)."
    )

    table = doc.add_table(rows=5, cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Window ending"
    hdr[1].text = "Value (€B)"
    hdr[2].text = "YoY value"
    hdr[3].text = "YoY kg"
    rows_data = [
        ("2025-11-01", "24.8", "+38.2%", "+71.1%"),
        ("2025-12-01", "25.9", "+36.5%", "+70.2%"),
        ("2026-01-01", "26.6", "+35.1%", "+69.8%"),
        ("2026-02-01", "27.3", "+34.5%", "+69.4%"),
    ]
    for r, row in enumerate(rows_data, start=1):
        cells = table.rows[r].cells
        for c, val in enumerate(row):
            cells[c].text = val
    # Bold the header row
    for cell in hdr:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True

    # ---- Featured finding 2: grouped bar chart ----
    doc.add_heading("Finished cars (broad) — EU-27 exports to China", level=1)

    doc.add_paragraph(
        "EU-27 exports of finished cars to China fell 40.7% YoY in value "
        "and 38.4% in volume — the largest decline among the cycle's Top 5 "
        "movers. Germany and France between them account for roughly 80% "
        "of the loss; both reporters' bilateral series show consistent "
        "negative YoYs across the last six monthly windows."
    )

    bar_png = grouped_bar_png(
        ["Germany", "France", "Italy", "Spain", "Belgium"],
        [4.21, 1.92, 0.85, 0.41, 0.29],
        [2.48, 1.13, 0.51, 0.24, 0.19],
        series_a_name="Prior 12mo (€B)",
        series_b_name="Current 12mo (€B)",
        title="EU-27 exports of finished cars to China — top-5 reporter breakdown",
        ylabel="€ billions",
    )
    doc.add_picture(io.BytesIO(bar_png), width=Mm(190))

    # ---- A small caveat / footnote block ----
    doc.add_heading("Caveats and provenance", level=2)
    caveats = [
        "Values are EU-27 reporter-side, Eurostat COMEXT v2 (CIF for imports, FOB for exports).",
        "EV battery figures cover HS 8507.60 only; broader Li-ion category (8507.6020/30) gives a slightly larger total.",
        "Prior-window values reflect the post-184-fix figures (unit-field bug corrected 2026-05-14).",
        "Source obs_ids 12345–12378 (current), 11890–11923 (prior). See 04_Data.xlsx for the full row set.",
    ]
    for c in caveats:
        doc.add_paragraph(c, style="List Bullet")

    doc.save(out_path)


# ---------------------------------------------------------------------------
# Xlsx generation — native LineChart object (the thing we want to survive
# upload-to-Sheets as an editable native chart, not a flat image)
# ---------------------------------------------------------------------------

def build_xlsx(out_path: Path) -> None:
    wb = Workbook()

    # ---- Sheet 1: EV batteries time series with line chart ----
    ws1 = wb.active
    ws1.title = "EV batteries (Li-ion)"

    ws1["A1"] = "Month"
    ws1["B1"] = "Prior 12mo (€B)"
    ws1["C1"] = "Current 12mo (€B)"
    for cell in ("A1", "B1", "C1"):
        ws1[cell].font = Font(bold=True)

    for i, (month, prior, current) in enumerate(zip(MONTHS, EV_BATTERY_PRIOR, EV_BATTERY_CURRENT), start=2):
        ws1.cell(row=i, column=1, value=month)
        ws1.cell(row=i, column=2, value=prior)
        ws1.cell(row=i, column=3, value=current)

    chart1 = LineChart()
    chart1.title = "EU-27 imports of Li-ion EV batteries from China"
    chart1.y_axis.title = "€ billions"
    chart1.x_axis.title = "Month"
    chart1.height = 9
    chart1.width = 18

    data = Reference(ws1, min_col=2, max_col=3, min_row=1, max_row=len(MONTHS) + 1)
    categories = Reference(ws1, min_col=1, min_row=2, max_row=len(MONTHS) + 1)
    chart1.add_data(data, titles_from_data=True)
    chart1.set_categories(categories)

    ws1.add_chart(chart1, "E2")

    # ---- Sheet 2: Finished cars bar chart ----
    ws2 = wb.create_sheet("Finished cars")
    ws2["A1"] = "Reporter"
    ws2["B1"] = "Prior 12mo (€B)"
    ws2["C1"] = "Current 12mo (€B)"
    for cell in ("A1", "B1", "C1"):
        ws2[cell].font = Font(bold=True)

    bar_rows = [
        ("Germany", 4.21, 2.48),
        ("France", 1.92, 1.13),
        ("Italy", 0.85, 0.51),
        ("Spain", 0.41, 0.24),
        ("Belgium", 0.29, 0.19),
    ]
    for i, (reporter, prior, current) in enumerate(bar_rows, start=2):
        ws2.cell(row=i, column=1, value=reporter)
        ws2.cell(row=i, column=2, value=prior)
        ws2.cell(row=i, column=3, value=current)

    chart2 = BarChart()
    chart2.type = "col"
    chart2.grouping = "clustered"
    chart2.title = "EU-27 exports of finished cars to China — top-5 reporter breakdown"
    chart2.y_axis.title = "€ billions"
    chart2.height = 9
    chart2.width = 18

    data2 = Reference(ws2, min_col=2, max_col=3, min_row=1, max_row=len(bar_rows) + 1)
    categories2 = Reference(ws2, min_col=1, min_row=2, max_row=len(bar_rows) + 1)
    chart2.add_data(data2, titles_from_data=True)
    chart2.set_categories(categories2)

    ws2.add_chart(chart2, "E2")

    wb.save(out_path)


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    docx_path = OUT_DIR / "03_Findings_test.docx"
    xlsx_path = OUT_DIR / "04_Data_test.xlsx"

    build_docx(docx_path)
    build_xlsx(xlsx_path)

    print(f"Wrote: {docx_path}  ({docx_path.stat().st_size / 1024:.1f} KB)")
    print(f"Wrote: {xlsx_path}  ({xlsx_path.stat().st_size / 1024:.1f} KB)")
    print()
    print("Next: drag both files to a Drive folder, double-click each, choose")
    print("'Open with Google Docs' / 'Open with Google Sheets'. Eyeball the")
    print("result against the local files. Things to look for are listed at")
    print("the top of dev_notes/2026-05-16_docx-drive-spike.md.")


if __name__ == "__main__":
    main()
