"""Markdown → docx translator for the briefing pack.

Walks a mistune AST of a rendered findings.md document and emits the
equivalent shape into a python-docx Document. Charts can be injected
at specific anchor points (currently after each top-mover list item)
via a callable passed in at construction.

Why not pandoc:
- meridian is pure-Python with no subprocess dependencies; pandoc
  would add an external binary requirement.
- Our markdown subset is small (~10 block types, ~5 inline types).
  A focused translator stays under 400 lines and gives full output
  control.

Why this exists as a separate module from `briefing_pack.docx`:
- `briefing_pack.docx` orchestrates: page setup, top-movers chart
  data fetch, final write.
- `briefing_pack.md_to_docx` is a pure translator: in = markdown
  string, out = blocks written to a Document. Reusable for any
  future docx-from-markdown surface (leads.docx, groups.docx, etc.).

The mistune AST node types we handle (and what we don't):

Handled:
- heading (levels 1-6)
- paragraph
- list (ordered + unordered, including nested)
- list_item, block_text
- table, table_head, table_body, table_row, table_cell
- block_quote
- block_code (fenced)
- thematic_break (horizontal rule)
- text, strong, emphasis, codespan, link
- softbreak, linebreak

Not handled / passed through as plain text:
- HTML blocks / inline HTML (none emitted by our sections today)
- Footnotes, definition lists, task lists (not in our markdown subset)
- Images (sections emit none; provenance bundle is link-only)
"""

from __future__ import annotations

import logging
import re
from typing import Callable

import mistune
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor

log = logging.getLogger(__name__)


# Finding-token regex shared with briefing_pack/render.py's
# _editorially_fresh_finding_ids. Lifted here so the translator can
# detect chart-injection anchors in list items.
_FINDING_TOKEN_RE = re.compile(r"\bfinding/(\d+)\b")


# Type alias: callable that takes a finding id and returns PNG bytes
# (or None if no chart is available for that finding).
ChartLookup = Callable[[int], bytes | None]


class MarkdownToDocxTranslator:
    """Walk a mistune AST and write each block into a python-docx
    Document. Inline runs (bold / italic / code / link) become run
    objects with the right styling inside the appropriate block.

    Chart injection: when `chart_for_finding` is provided, after each
    list item whose text carries a `finding/{id}` token, the
    translator inserts a chart picture if `chart_for_finding(id)`
    returns PNG bytes.

    Stateless apart from the target Document and the chart-lookup
    callable. Safe to instantiate per-render.
    """

    def __init__(
        self,
        doc: Document,
        *,
        chart_for_finding: ChartLookup | None = None,
        chart_width_mm: int = 190,
    ) -> None:
        self.doc = doc
        self.chart_for_finding = chart_for_finding
        self.chart_width_mm = chart_width_mm
        # mistune parser with table support (the briefing pack's tables
        # use the GFM pipe syntax).
        self._parse = mistune.create_markdown(
            renderer=None, plugins=["table", "strikethrough"],
        )

    # ----- public entry point ------------------------------------------------

    def translate(self, markdown_text: str) -> None:
        """Translate `markdown_text` into the underlying Document.

        Iterates the top-level AST nodes; each block emits one or more
        docx blocks. Existing Document content is preserved (the
        translator appends).
        """
        ast = self._parse(markdown_text)
        if not isinstance(ast, list):
            # Defensive — mistune always returns a list for the top level,
            # but renderer=None can return other shapes in edge cases.
            log.warning("Unexpected AST root: %r — skipping", type(ast))
            return
        for node in ast:
            self._handle_block(node)

    # ----- block dispatch ----------------------------------------------------

    def _handle_block(self, node: dict) -> None:
        t = node.get("type", "")
        handler = getattr(self, f"_block_{t}", None)
        if handler is None:
            # Unknown block — pass through as plain paragraph if we can
            # extract any text. Real-world cases: future markdown features
            # we haven't taught the translator yet.
            text = self._collect_plain_text(node)
            if text.strip():
                self.doc.add_paragraph(text)
            return
        handler(node)

    # ----- block handlers ---------------------------------------------------

    def _block_blank_line(self, node: dict) -> None:
        # mistune emits these between blocks; python-docx handles
        # paragraph spacing via styles, so we drop them.
        pass

    def _block_heading(self, node: dict) -> None:
        level = (node.get("attrs") or {}).get("level", 1)
        # python-docx: level 0 is Title, 1-9 are Heading 1-9
        # Cap at level 9 so we don't error on absurd inputs.
        docx_level = max(0, min(9, level))
        heading_p = self.doc.add_heading(level=docx_level)
        self._render_inline_into_paragraph(node.get("children") or [], heading_p)

    def _block_paragraph(self, node: dict) -> None:
        p = self.doc.add_paragraph()
        self._render_inline_into_paragraph(node.get("children") or [], p)

    def _block_thematic_break(self, node: dict) -> None:
        # Horizontal rule. python-docx doesn't have a first-class API
        # for this; emit a bottom-bordered empty paragraph.
        p = self.doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "auto")
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _block_table(self, node: dict) -> None:
        """GFM pipe table. mistune emits `table` with `table_head` and
        `table_body` children. The head's children are `table_cell`
        directly; the body's are `table_row` wrappers around cells.

        Output: a python-docx Table with the Light Grid style, bold
        header row, and inline content per cell preserving bold /
        italic / link runs. Column-align attrs are honoured via the
        paragraph's alignment.
        """
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        # Collect rows: first the header (one row), then body rows.
        header_cells: list[dict] = []
        body_rows: list[list[dict]] = []
        for child in node.get("children") or []:
            ctype = child.get("type")
            if ctype == "table_head":
                # Head children are table_cell nodes directly.
                header_cells = [
                    c for c in (child.get("children") or [])
                    if c.get("type") == "table_cell"
                ]
            elif ctype == "table_body":
                for row in child.get("children") or []:
                    if row.get("type") != "table_row":
                        continue
                    body_rows.append([
                        c for c in (row.get("children") or [])
                        if c.get("type") == "table_cell"
                    ])

        n_cols = max(
            len(header_cells),
            max((len(r) for r in body_rows), default=0),
        )
        if n_cols == 0:
            return

        table = self.doc.add_table(rows=1 + len(body_rows), cols=n_cols)
        try:
            table.style = "Light Grid Accent 1"
        except KeyError:
            # Style not available in some docx templates — fall back to
            # whatever the default is; the table still renders.
            pass

        align_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
        }

        # Header
        for col_idx, cell_node in enumerate(header_cells):
            if col_idx >= n_cols:
                break
            cell = table.rows[0].cells[col_idx]
            # Replace the auto-inserted empty paragraph
            cell.text = ""
            p = cell.paragraphs[0]
            align = (cell_node.get("attrs") or {}).get("align")
            if align in align_map:
                p.alignment = align_map[align]
            self._render_inline_into_paragraph(
                cell_node.get("children") or [], p, force_bold=True,
            )

        # Body
        for row_idx, row_cells in enumerate(body_rows):
            docx_row = table.rows[1 + row_idx]
            for col_idx, cell_node in enumerate(row_cells):
                if col_idx >= n_cols:
                    break
                cell = docx_row.cells[col_idx]
                cell.text = ""
                p = cell.paragraphs[0]
                align = (cell_node.get("attrs") or {}).get("align")
                if align in align_map:
                    p.alignment = align_map[align]
                self._render_inline_into_paragraph(
                    cell_node.get("children") or [], p,
                )

    def _block_block_code(self, node: dict) -> None:
        """Fenced code block. Renders as a paragraph in a monospace
        font; preserves leading whitespace by using soft breaks for
        embedded newlines."""
        text = node.get("raw") or ""
        p = self.doc.add_paragraph()
        for i, line in enumerate(text.splitlines()):
            if i > 0:
                p.add_run().add_break()
            run = p.add_run(line)
            run.font.name = "Courier New"
            run.font.size = Pt(9)

    def _block_block_quote(self, node: dict) -> None:
        """Block quote — each contained block emitted as an italic
        indented paragraph. Nested blocks (rare) flatten into the
        same indent level."""
        for child in node.get("children") or []:
            if child.get("type") == "paragraph":
                p = self.doc.add_paragraph()
                p.paragraph_format.left_indent = Pt(18)
                # Inline children emitted with italic flag forced.
                self._render_inline_into_paragraph(
                    child.get("children") or [], p, force_italic=True,
                )
            else:
                # Anything other than a paragraph inside a quote — pass
                # through plain text as a fallback.
                text = self._collect_plain_text(child).strip()
                if text:
                    p = self.doc.add_paragraph(text)
                    p.paragraph_format.left_indent = Pt(18)
                    for r in p.runs:
                        r.italic = True

    def _block_list(self, node: dict) -> None:
        """Ordered or unordered list. Iterates list_item children and
        emits each as a styled paragraph. Nested lists are handled
        recursively (each nested list emits at the same flat level —
        python-docx's first-class list nesting requires deeper XML
        plumbing than this v1 needs)."""
        attrs = node.get("attrs") or {}
        ordered = attrs.get("ordered", False)
        style_name = "List Number" if ordered else "List Bullet"
        for child in node.get("children") or []:
            if child.get("type") != "list_item":
                continue
            self._render_list_item(child, style_name=style_name)

    def _render_list_item(self, node: dict, *, style_name: str) -> None:
        """Emit a list_item as a list-styled paragraph. The item's
        block_text holds the inline content; nested lists become
        subsequent paragraphs (flat for v1)."""
        for sub in node.get("children") or []:
            t = sub.get("type")
            if t == "block_text":
                p = self.doc.add_paragraph(style=style_name)
                self._render_inline_into_paragraph(
                    sub.get("children") or [], p,
                )
                # After writing the item's prose, look for a finding/N
                # token and inject a chart if the lookup provides one.
                self._maybe_inject_chart_after_item(sub)
            elif t == "paragraph":
                # Multi-paragraph list item — second paragraph onward
                # gets the same list style.
                p = self.doc.add_paragraph(style=style_name)
                self._render_inline_into_paragraph(
                    sub.get("children") or [], p,
                )
            elif t == "list":
                # Nested list — recurse. python-docx doesn't carry
                # depth automatically; the nested list paragraphs use
                # the same style. Visual fidelity is acceptable for our
                # use case (single-level nesting in the existing pack).
                self._block_list(sub)
            else:
                # Unknown sub-element — fall back to plain text.
                text = self._collect_plain_text(sub).strip()
                if text:
                    self.doc.add_paragraph(text, style=style_name)

    def _maybe_inject_chart_after_item(self, block_text_node: dict) -> None:
        """Scan the item's text for a finding/N token; if found, ask
        the chart-lookup callable for PNG bytes and insert as a
        picture in the current paragraph flow.

        Chart sits on its own paragraph immediately after the list
        item's prose paragraph, which the docx renderer for top-N
        movers expects."""
        if self.chart_for_finding is None:
            return
        text = self._collect_plain_text(block_text_node)
        m = _FINDING_TOKEN_RE.search(text)
        if m is None:
            return
        finding_id = int(m.group(1))
        png = self.chart_for_finding(finding_id)
        if not png:
            return
        import io
        from docx.shared import Mm
        self.doc.add_picture(
            io.BytesIO(png), width=Mm(self.chart_width_mm),
        )

    # ----- inline rendering --------------------------------------------------

    def _render_inline_into_paragraph(
        self, children: list[dict], paragraph, *,
        force_italic: bool = False,
        force_bold: bool = False,
    ) -> None:
        """Walk inline AST nodes and add corresponding runs to the
        target paragraph. Style flags compose: a `<strong>` inside a
        block_quote (force_italic) yields bold+italic runs, matching
        the markdown's nested emphasis semantics."""
        for child in children:
            self._render_inline_node(
                child, paragraph,
                bold=force_bold, italic=force_italic, code=False,
                link_url=None,
            )

    def _render_inline_node(
        self, node: dict, paragraph, *,
        bold: bool, italic: bool, code: bool, link_url: str | None,
    ) -> None:
        t = node.get("type", "")
        if t == "text":
            text = node.get("raw") or ""
            if not text:
                return
            if link_url:
                self._add_hyperlink_run(
                    paragraph, text, link_url,
                    bold=bold, italic=italic, code=code,
                )
            else:
                run = paragraph.add_run(text)
                run.bold = bold
                run.italic = italic
                if code:
                    run.font.name = "Courier New"
        elif t == "strong":
            for sub in node.get("children") or []:
                self._render_inline_node(
                    sub, paragraph,
                    bold=True, italic=italic, code=code, link_url=link_url,
                )
        elif t == "emphasis":
            for sub in node.get("children") or []:
                self._render_inline_node(
                    sub, paragraph,
                    bold=bold, italic=True, code=code, link_url=link_url,
                )
        elif t == "codespan":
            text = node.get("raw") or ""
            run = paragraph.add_run(text)
            run.font.name = "Courier New"
            run.bold = bold
            run.italic = italic
        elif t == "link":
            url = (node.get("attrs") or {}).get("url", "")
            for sub in node.get("children") or []:
                self._render_inline_node(
                    sub, paragraph,
                    bold=bold, italic=italic, code=code, link_url=url,
                )
        elif t == "softbreak":
            paragraph.add_run(" ")
        elif t == "linebreak":
            paragraph.add_run().add_break()
        elif t == "strikethrough":
            for sub in node.get("children") or []:
                run_p = paragraph.add_run()
                # Apply strikethrough via XML — not exposed as a
                # first-class python-docx attribute on Font.
                rPr = run_p._element.get_or_add_rPr()
                strike = OxmlElement("w:strike")
                strike.set(qn("w:val"), "true")
                rPr.append(strike)
                # Render the sub-content into this same paragraph;
                # post-style applied above.
                self._render_inline_node(
                    sub, paragraph,
                    bold=bold, italic=italic, code=code, link_url=link_url,
                )
        else:
            # Unknown inline — fall back to its plain text. Keeps
            # forward compatibility if mistune introduces new node
            # types or sections start emitting features we haven't
            # taught the translator.
            text = self._collect_plain_text(node)
            if text:
                run = paragraph.add_run(text)
                run.bold = bold
                run.italic = italic
                if code:
                    run.font.name = "Courier New"

    def _add_hyperlink_run(
        self, paragraph, text: str, url: str, *,
        bold: bool, italic: bool, code: bool,
    ) -> None:
        """Insert a clickable hyperlink run into a paragraph.

        python-docx lacks a first-class hyperlink API, so we construct
        the underlying OOXML directly: a `w:hyperlink` element wrapping
        a `w:r` run, with the relationship stored in the document's
        relationships table.
        """
        part = paragraph.part
        r_id = part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)

        run_el = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        # Apply hyperlink style + colour + underline
        rStyle = OxmlElement("w:rStyle")
        rStyle.set(qn("w:val"), "Hyperlink")
        rPr.append(rStyle)
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "0563C1")
        rPr.append(color)
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)
        if bold:
            b = OxmlElement("w:b")
            rPr.append(b)
        if italic:
            i = OxmlElement("w:i")
            rPr.append(i)
        if code:
            rFonts = OxmlElement("w:rFonts")
            rFonts.set(qn("w:ascii"), "Courier New")
            rFonts.set(qn("w:hAnsi"), "Courier New")
            rPr.append(rFonts)
        run_el.append(rPr)

        text_el = OxmlElement("w:t")
        text_el.set(qn("xml:space"), "preserve")
        text_el.text = text
        run_el.append(text_el)

        hyperlink.append(run_el)
        paragraph._p.append(hyperlink)

    # ----- helpers ----------------------------------------------------------

    def _collect_plain_text(self, node: dict) -> str:
        """Recursively concatenate `raw` from every text-like descendant.
        Used as a fallback for unknown block types and for the
        finding-token scan in list items."""
        if "raw" in node and "children" not in node:
            return node.get("raw") or ""
        parts: list[str] = []
        for child in node.get("children") or []:
            parts.append(self._collect_plain_text(child))
        return "".join(parts)
